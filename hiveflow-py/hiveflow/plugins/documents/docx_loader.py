"""DOCX document loader using python-docx."""

from pathlib import Path

from hiveflow.plugins.documents import Document, DocumentLoaderPlugin


class DocxLoader(DocumentLoaderPlugin):
    """Loader for DOCX files using python-docx."""

    @property
    def plugin_id(self) -> str:
        return "docx"

    @property
    def description(self) -> str:
        return "DOCX file loader using python-docx"

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def load(self, file_path: str | Path) -> Document:
        try:
            import docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX support. Install with: pip install python-docx"
            ) from exc

        path = Path(file_path)
        doc = docx.Document(str(path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Prefix headings with markdown-style markers
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split()[-1])
                    parts.append(f"{'#' * level} {text}")
                except (ValueError, IndexError):
                    parts.append(text)
            else:
                parts.append(text)

        content = "\n\n".join(parts)
        return Document(
            content=content,
            source=str(path),
            metadata={
                "filename": path.name,
                "size_bytes": path.stat().st_size,
                "paragraph_count": len(parts),
            },
        )
